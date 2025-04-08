import os
import logging
import yaml
import PyPDF2
import pandas as pd
from langchain.document_loaders import TextLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import SentenceTransformerEmbeddings
from langchain.schema import Document

# Setup logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

# Load configurations
def load_config(config_path="config.yaml"):
    with open(config_path, "r") as f:
        return yaml.safe_load(f)

config = load_config()
CHUNK_SIZE = config['chunk_size']
CHUNK_OVERLAP = config['chunk_overlap']
PERSIST_DIRECTORY = config['persist_directory']
EMBEDDING_MODEL = config.get('embedding_model', "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")
file_metadata = {}

# Function to load a document based on its file extension
def load_file(filepath):
    try:
        logger.info(f"Loading file: {filepath}")
        if filepath.endswith('.txt'):
            loader = TextLoader(filepath, encoding='utf-8')
            return loader.load()
        elif filepath.endswith('.pdf'):
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = "".join(page.extract_text() for page in reader.pages)
            return [Document(page_content=text)]
        elif filepath.endswith('.xlsx') or filepath.endswith('.xls'):
            # For Excel files, we'll convert to markdown table format
            df = pd.read_excel(filepath)
            markdown_text = df.to_markdown()
            return [Document(page_content=markdown_text)]
        elif filepath.endswith('.docx'):
            try:
                import docx
                doc = docx.Document(filepath)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                return [Document(page_content=text)]
            except ImportError:
                logger.warning("python-docx not installed, skipping docx file")
                return []
        elif filepath.endswith('.md'):
            with open(filepath, 'r', encoding='utf-8') as file:
                text = file.read()
            return [Document(page_content=text)]
        else:
            logger.warning(f"Unsupported file type: {filepath}")
    except Exception as e:
        logger.error(f"Failed to load {filepath}: {e}")
    return []

# Initialize or load the vector database
def initialize_vector_database():
    if os.path.exists(PERSIST_DIRECTORY):
        logger.info("Loading vector database from persisted state.")
        return Chroma(persist_directory=PERSIST_DIRECTORY, 
                      embedding_function=SentenceTransformerEmbeddings(model_name=EMBEDDING_MODEL))
    else:
        logger.info("No existing vector database found. Initializing a new one.")
        return Chroma(embedding_function=SentenceTransformerEmbeddings(model_name=EMBEDDING_MODEL), 
                      persist_directory=PERSIST_DIRECTORY)

# Initialize the vector database
vector_db = initialize_vector_database()

def save_metadata():
    """Save file metadata to config.yaml to persist across restarts."""
    try:
        if os.path.exists("config.yaml"):
            with open("config.yaml", "r") as f:
                config = yaml.safe_load(f) or {}
        else:
            config = {}

        # Convert sets to lists before saving
        config["file_metadata"] = {key: list(value) for key, value in file_metadata.items()}

        with open("config.yaml", "w") as f:
            yaml.safe_dump(config, f)

        logger.info("File metadata saved successfully.")
    except Exception as e:
        logger.error(f"Error saving file metadata: {e}")

def load_metadata():
    """Load file metadata from config.yaml and convert lists back to sets."""
    global file_metadata
    try:
        if os.path.exists("config.yaml"):
            with open("config.yaml", "r") as f:
                config = yaml.safe_load(f)
                if "file_metadata" in config:
                    file_metadata = {key: set(value) for key, value in config.get("file_metadata", {}).items()}
                    logger.info("File metadata loaded successfully.")
        else:
            file_metadata = {}
    except Exception as e:
        logger.error(f"Error loading file metadata: {e}")
        file_metadata = {}

# Function to process and add documents to the vector database
def process_and_add_documents(filepaths, additional_metadata=None):
    try:
        documents = []
        filepath_to_docs = {}  # Track which documents came from which files
        
        for filepath in filepaths:
            docs = load_file(filepath)
            
            # Add metadata to documents
            if additional_metadata:
                for doc in docs:
                    doc.metadata.update(additional_metadata)
                    
            # Add filepath to metadata
            for doc in docs:
                doc.metadata["source"] = filepath
                doc.metadata["filename"] = os.path.basename(filepath)
            
            documents.extend(docs)
            filepath_to_docs[filepath] = len(docs)  # Store how many docs came from this file
        
        # Split documents into chunks
        text_splitter = CharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
        chunks = text_splitter.split_documents(documents)
        
        # Add chunks to vector database
        chunk_ids = vector_db.add_documents(chunks)
        vector_db.persist()
        
        # Update file metadata - track which file each chunk came from
        # For simplicity, we'll attribute each chunk to its source file
        # This is a simplified approach - in reality you may need document tracking
        for i, chunk_id in enumerate(chunk_ids):
            # Find which file this chunk came from based on the chunk's metadata
            source = chunks[i].metadata.get("source")
            if source:
                filepath = source
                filename = os.path.basename(filepath)
                if filename not in file_metadata:
                    file_metadata[filename] = set()
                file_metadata[filename].add(chunk_id)
        
        save_metadata()
        logger.info(f"Processed and added {len(chunks)} document chunks.")
        return True
    except Exception as e:
        logger.error(f"Error processing documents: {e}")
        raise